# -*- coding: utf-8 -*-
"""
Created on Fri Feb  9 16:14:24 2018

@author: Shivank.r
"""
from __future__ import (
    absolute_import, division, print_function, unicode_literals
)
import Initial_processing
import re
import pandas as pd
import variable_codes
from docx import Document
from docx.table import  Table
from docx.text.paragraph import Paragraph
#from dateutil import parser


class Docx_class:
    def __init__(self,filepath):
        self.keyset={'employer','languages', 'forte', 'intrests', 'skiills', 'personnel', 'i', 'engineering', 'technical', 'detail','details', 'awards', 'received', 'sets', 'computer', 'occupational', 'details.', 'qualifying', 'highlights', 'profeciency', 'primary', 'qualicaion', 'competenancy', 'qualifications', 'particulars', 'objecrive', 'p', 'purview', 'snippets', 'educations', 'techinal', 'technicall', 'competency', 'in', '&', 'expertise', 'development', 'management', 'background', 'soft', 'academics', 'career', 'techanical',  'core', 'compitencies', 'scholasitics', 'employee', 'key', 'history',  'experience', 'scholastics', 'competencies', 'dossier', 'experiences', 'detail', 'academia', 'organization', 'proofessional', 'acquired',  'experience.', 'niceties', 'certifications&', 'certifications', 'strengths', 'set', 'education', 'peofessional', 'designations', 'qualications', 'personalprofile', 'certification', 'experiances', 'abridgement', 'info', 'chronicle', 'journey', 'exposures', 'rating', '&certifications', 'skillset', 'record', 'tools', 'domain', 'skillsets', 'deatils', 'overview', 'techncial', 'academics', 'features', 'organisational',  'capabilities',  'prsonal', 'known', 'knowledge', 'informations', 'h', 'contour', 'acadamic', 'peronal', 'courses', 'professional', 'technologies', 'summery', 'goal', 'used', 'employement', 'employment', 'i.t.',  'higher', 'professionalexperience', 'training', 'carrer', 'job', 'profiles', 'back', 'objective', 'skills&', 'techinical', 'achivement', 'techbical',  'profile', 'educatioin', 'proffesional', 'scan', 'experince', 'ployement', 'chronicles', 'expertice', 'experiance', 'minutiae', 'profesional', 'competence', 'scholastic',  'acedemic', 'business', 'records', 'industrial', 'skils', 'responsibilities', 'quallification', 'inventory',  'trainings', 'experiecne', 'perssional', 'technicalskill', 'accolades', 'synopsys',  'professsional', 'qulification', 'ground', 'acdemic', 'acadmic', 'profiency', 'objetive', 'empoyment', 'workin', 'technology', 'competences', 'persional', 'personal', 'skill', 'dossiere', 'certificate', 'synopsis', 'oualification', 'scholatics', 'porfessional', 'strength', 'information', 'exposure', 'expierence', 'cerfication', 'software', 'objectives', 'sill', 'history', 'path', 'work', 'recital', 'industry', 'summary', 'possessed', 'initial', 'qualifactions', 'chronology', 'abridgemen', 'it', 'qualifiaction', 'preview', 'educationational', 'category', 'n', 'careers', 'detailed', 'abilities', 'snapshot',  'techinacal',  'educational', 'professinol', 'credentials', 'database', 'qulifications', 'educatonal', 'technicle', 'achievement', 'my', 'outline', 'expereince', 'employers', 'organisation', 'technicals', 'skills', 'functional', 'worked', 'qualification', 'dossier', 'exeperience', 'experirnce', 'its', 'proficiencies', 'areas', 'interests', 'credential', 'data', 'corporate', 'progression', 'proficiency', 'skilss', 'edcuational', 'expericence', 'techincal', 'experinece', 'skills.', 'history', 'analytics', 'matrix', 's', 'organizational', 'm', 'accomplishments', 'academic', 'qualificaion', 'skills', 'exprience', 'rofessional', 'eductional', 'acadamicqualification', 'highlight', 'other', 'skills&expierence', 'carrier', 'information', 'relevant'}
        self.document=Document(filepath)
        self.keys=pd.read_excel('/home/shivank/parser/parser_class/all_resumes_keywords_final.xlsx',encoding = 'iso-8859-1')
        self.array=Initial_processing.items(self.document)
        self.resume_string = Initial_processing.docs_conversion(self.array)
        self.relations=Initial_processing.relations(self.resume_string)
        self.relations=[[k.lower().strip(),v.strip()] for k,v in self.relations ]
        self.tags=[]
        self.tagged_data=[]
        self.single_cell_tranformation()
        self.table_array=list(filter(lambda x:isinstance(x, Table),self.array))
        self.table_tagger()
        self.keywords=[]
        self.tagged_para=[]
        self.imp_keywords()
        self.bifurcation()
        self.tagged_relations()
        self.email=None
        self.hyperlink_email()
        self.keyword_tag_dict = {'Experience':{'tagged_relations':[],'tagged_para':[],'tagged_data': []},'Skills':{'tagged_relations':[],'tagged_para':[],'tagged_data': []},'Mobile' : {'tagged_relations':[],'tagged_para':[],'tagged_data': []},'Education':{'tagged_relations':[],'tagged_para':[],'tagged_data': []},'DOB':{'tagged_relations':[],'tagged_para':[],'tagged_data': []},'Name':{'tagged_relations':[],'tagged_para':[],'tagged_data': []},'Location' : {'tagged_relations':[],'tagged_para':[],'tagged_data': []} ,'Email': {'tagged_relations':[],'tagged_para':[],'tagged_data': []} ,'Gender' : {'tagged_relations':[],'tagged_para':[],'tagged_data': []}}
        self.keyword_tag_extract()
        
    def single_cell_tranformation(self):
        for i,element in enumerate(self.array):
            if isinstance(element, Table):
                if len(element.rows)==1 and len(element.columns)==1:
                    element.rows[0].cells[0].paragraphs[0].runs[0].font.bold=True
                    self.array[i]=element.rows[0].cells[0].paragraphs[0]
    
    def table_tagger(self):
        for table in self.table_array:
            data=[]
            for i, row in enumerate(table.rows):
                text = list(cell.text.strip() for cell in row.cells)
                data.append(text)
            ind=self.array.index(table)
            if len(data)==2 and len(set(data[0]))==1:
                self.tagged_data.append([data[0][0],data[1:]])
                self.tags.append(data[0][0])
            if len(data)>=2:
                if data[0].count(data[0][0])==len(data[0]):
                    self.tagged_data.append([data[0][0],data[1:]])
                    self.tags.append(data[0][0])
                    continue
                while(ind>0):
                    ind=ind-1
                    if isinstance(self.array[ind], Paragraph):
                        if all(list(map(lambda y:y.bold,self.array[ind].runs))) and (self.array[ind].text.strip().istitle() or self.array[ind].text.strip().isupper()) and re.match('^[A-z ]+:?$',self.array[ind].text.strip()):
                            self.tagged_data.append([self.array[ind].text.strip(),data])
                            self.tags.append(self.array[ind].text.strip())
                            break
        
    def imp_keywords(self):
        for i in self.array:
            if isinstance(i, Paragraph):
                if re.sub("[\:\-\'\t\\\/]","",i.text.lower().strip()) in self.keys[self.keys["Heading"]=="Header"]["keyword"].tolist():
                    self.keywords.append(i.text.strip())
                    continue
                if all(list(map(lambda y:y.bold,i.runs))) and (i.text.strip().istitle() or i.text.strip().isupper() )and re.match('[A-z &\/\\\.\,\-]+:?$',i.text.strip())and len(set(re.sub("[\:\-\'\t\\\/]","",i.text.lower().strip()).split()).intersection(self.keyset))>0:
                    if (i.text.strip()) not in list(self.keys[self.keys['Heading'] == 'Remove']['keyword']):        
                        self.keywords.append(i.text.strip())   
    
    def bifurcation(self):
            text=[j.text.strip() for j in list(filter(lambda x:isinstance(x,Paragraph),self.array))]
            for i,key in enumerate(self.keywords):
                if key in self.tags:
                    continue
                if i==len(self.keywords)-1:
                    self.tagged_para.append([key,text[text.index(key)+1:]])
                else:
                   self.tagged_para.append([key,text[text.index(key)+1:text.index(self.keywords[i+1])]]) 

    def tagged_relations(self):
        keywds = pd.read_csv('/home/shivank/parser/parser_class/relation_keywords.csv',encoding="iso-8859-1")
        keywds = keywds[keywds['keyword'].isin(list(map(lambda x:x[0],self.relations)))][keywds['keyword tag'] != 'Remove']
        keywds['value'] = keywds['keyword'].apply(lambda x: [y[1]  for y in self.relations if y[0].lower()==x])
        keywds['value'][keywds['keyword tag']!="Skills"]=keywds['value'][keywds['keyword tag']!="Skills"].apply(lambda x:x[0])
        keywds['value'][keywds['keyword tag']=="Skills"]=keywds['value'][keywds['keyword tag']=="Skills"].apply(lambda x:','.join(x))
        self.relations={keywds['keyword tag'][i]:keywds['value'][i] for i in  keywds.index}
    def hyperlink_email(self):
        rels = self.document.part.rels
        email=[]
        for rel in Initial_processing.iter_hyperlink_rels(rels):
            email.append(re.sub("mailto:","",rels[rel]._target))
        self.email="/".join(email)
           
    def keyword_tag_extract(self):
        self.keywords_dict = self.keys[self.keys['keyword tag'] != 'Remove'].groupby('keyword tag')['keyword'].apply(list).to_dict()
        for i in self.tagged_para:
            if re.sub("[\:\-\'\t\\\/]","",i[0].lower().strip()) in self.keywords_dict['Experience']:
                self.keyword_tag_dict['Experience']['tagged_para'].append(i[1])
            elif re.sub("[\:\-\'\t\\\/]","",i[0].lower().strip()) in self.keywords_dict['Education']:
                self.keyword_tag_dict['Education']['tagged_para'].append(i[1])            
            elif re.sub("[\:\-\'\t\\\/]","",i[0].lower().strip()) in self.keywords_dict['Skills']:
                self.keyword_tag_dict['Skills']['tagged_para'].append(i[1])       
        
        for i in self.tagged_data:
            if re.sub("[\:\-\'\t\\\/]","",i[0].lower().strip()) in self.keywords_dict['Experience']:
                self.keyword_tag_dict['Experience']['tagged_data'].append(i[1])
            elif re.sub("[\:\-\'\t\\\/]","",i[0].lower().strip()) in self.keywords_dict['Education']:
                self.keyword_tag_dict['Education']['tagged_data'].append(i[1])        
            elif re.sub("[\:\-\'\t\\\/]","",i[0].lower().strip())in self.keywords_dict['Skills']:
                self.keyword_tag_dict['Skills']['tagged_data'].append(i[1])

        if 'Skills' in list(self.relations.keys()):
            for i in self.relations['Skills']:
                self.keyword_tag_dict['Skills']['tagged_relations'].append(i)               
        if 'DOB' in list(self.relations.keys()):
            self.keyword_tag_dict['DOB']['tagged_relations'].append(self.relations['DOB'])
        else:
            self.keyword_tag_dict['DOB']['tagged_relations'].append(variable_codes.dob(self.resume_string))
        if 'Mobile' in list(self.relations.keys()):
            self.keyword_tag_dict['Mobile']['tagged_relations'].append(self.relations['Mobile'])
        else:        
            self.keyword_tag_dict['Mobile']['tagged_relations'].append(variable_codes.mobile(self.resume_string))
        if self.email:
            self.keyword_tag_dict['Email']['tagged_relations'].append(variable_codes.email(self.email))
        elif 'Email' in list(self.relations.keys()):
            self.keyword_tag_dict['Email']['tagged_relations'].append(self.relations['Email'])
        else:
            self.keyword_tag_dict['Email']['tagged_relations'].append(variable_codes.email(self.resume_string))
        if 'Name' in list(self.relations.keys()):
            self.keyword_tag_dict['Name']['tagged_relations'].append(self.relations['Name'])
        else:
            self.keyword_tag_dict['Name']['tagged_relations'].append(variable_codes.names_extractor(self.resume_string))
        if 'Gender' in list(self.relations.keys()):
            self.keyword_tag_dict['Gender']['tagged_relations'].append(self.relations['Gender'])
#        else:
#           self.keyword_tag_dict['Gender']['tagged_relations'].append(variable_codes.gender_prediction(name))
        if 'Location' in list(self.relations.keys()):
            self.keyword_tag_dict['Location']['tagged_relations'].append(self.relations['Location'])
        else:
            self.keyword_tag_dict['Location']['tagged_relations'].append(variable_codes.location_parse(self.resume_string))

            
            
    
